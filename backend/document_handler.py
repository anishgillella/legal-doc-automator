import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Tuple, Optional


class DocumentHandler:
    def __init__(self, doc_path: str):
        """Initialize document handler with path to .docx file"""
        self.doc_path = doc_path
        self.doc = None
        self.full_text = ""
        self.paragraphs_data = []
        
    def load_document(self) -> bool:
        """Load the .docx document"""
        try:
            self.doc = Document(self.doc_path)
            self._extract_text_structure()
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False
    
    def _extract_text_structure(self):
        """Extract text while preserving paragraph structure and formatting info"""
        self.full_text = ""
        self.paragraphs_data = []
        
        # Extract from regular paragraphs
        for para_idx, para in enumerate(self.doc.paragraphs):
            para_text = para.text
            self.full_text += para_text + "\n"
            
            # Store paragraph metadata
            self.paragraphs_data.append({
                'index': para_idx,
                'text': para_text,
                'alignment': para.alignment,
                'style': para.style.name if para.style else 'Normal',
                'runs': [(run.text, {
                    'bold': run.bold,
                    'italic': run.italic,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                }) for run in para.runs]
            })
        
        # Extract from table cells
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    if cell_text.strip():
                        self.full_text += cell_text + "\n"
                        
                        # Store cell metadata (similar to paragraph)
                        self.paragraphs_data.append({
                            'index': f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                            'text': cell_text,
                            'alignment': None,
                            'style': 'Table Cell',
                            'runs': []
                        })
    
    def get_full_text(self) -> str:
        """Return extracted full text"""
        return self.full_text
    
    def get_paragraph_count(self) -> int:
        """Return number of paragraphs"""
        return len(self.doc.paragraphs)
    
    def replace_placeholder(self, placeholder: str, value: str) -> bool:
        """
        Simple replacement logic:
        - Explicit placeholder [text] → Replace entire placeholder with value
        - Blank field (ends with : or space) → Replace only the blank part, keep label
        
        Args:
            placeholder: The placeholder text to find
            value: The replacement value
        
        Returns:
            True if replacement was successful
        """
        try:
            replaced_count = 0
            
            # Determine type: explicit placeholder or blank field
            # Explicit: [text], {text}, (text), _____, etc - replace ENTIRE placeholder
            # Blank: ends with : or space - replace only the blank part, keep label
            is_explicit_placeholder = (
                (placeholder.startswith('[') and placeholder.endswith(']')) or
                (placeholder.startswith('{') and placeholder.endswith('}')) or
                (placeholder.startswith('(') and placeholder.endswith(')')) or
                '_' in placeholder  # Underscores are explicit
            )
            is_blank_field = placeholder.endswith(': ') or placeholder.endswith(':')
            
            # Build list of patterns to try (handle whitespace variations)
            patterns_to_try = [placeholder]
            if is_blank_field:
                base = placeholder.rstrip(': \t')  # Get the label without trailing space/colon
                patterns_to_try.extend([
                    base + ':\t',    # Tab variant
                    base + ':  ',    # Double space variant
                    base + ': ',     # Space variant
                    base + ': _____', # Underscores variant
                ])
            
            # Replace in paragraphs
            for para in self.doc.paragraphs:
                full_para_text = ''.join([run.text for run in para.runs])
                
                for pattern in patterns_to_try:
                    if pattern in full_para_text:
                        # SIMPLE RULE:
                        # - Explicit [placeholder]: replace entire placeholder
                        # - Blank field (ends with :): keep label, replace blank part
                        if is_explicit_placeholder:
                            new_text = full_para_text.replace(pattern, value, 1)
                        else:
                            # Blank field: find label position and keep it
                            label_pos = full_para_text.find(pattern)
                            if label_pos != -1:
                                new_text = full_para_text[:label_pos + len(pattern)] + value
                            else:
                                continue
                        
                        if new_text != full_para_text:
                            # Clear and rewrite paragraph
                            for run in para.runs:
                                r = run._element
                                r.getparent().remove(r)
                            
                            para.add_run(new_text)
                            replaced_count += 1
                            break  # Move to next paragraph
            
            # Replace in table cells
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_para_text = ''.join([run.text for run in para.runs])
                            
                            for pattern in patterns_to_try:
                                if pattern in full_para_text:
                                    if is_explicit_placeholder:
                                        new_text = full_para_text.replace(pattern, value, 1)
                                    else:
                                        label_pos = full_para_text.find(pattern)
                                        if label_pos != -1:
                                            new_text = full_para_text[:label_pos + len(pattern)] + value
                                        else:
                                            continue
                                    
                                    if new_text != full_para_text:
                                        for run in para.runs:
                                            r = run._element
                                            r.getparent().remove(r)
                                        
                                        para.add_run(new_text)
                                        replaced_count += 1
                                        break
            
            return replaced_count > 0
        except Exception as e:
            print(f"Error replacing placeholder: {e}")
            return False
    
    def replace_placeholder_at_position(self, placeholder: str, value: str, position_index: int = 0) -> bool:
        """
        Replace a specific occurrence (by position) of a placeholder.
        Uses simple logic: explicit [placeholder] or blank field.
        
        Args:
            placeholder: The placeholder text to find
            value: The replacement value  
            position_index: Which occurrence (0=first, 1=second, etc.)
        
        Returns:
            True if replacement was successful
        """
        try:
            # Determine type - handle all placeholder types
            is_explicit_placeholder = (
                (placeholder.startswith('[') and placeholder.endswith(']')) or
                (placeholder.startswith('{') and placeholder.endswith('}')) or
                (placeholder.startswith('(') and placeholder.endswith(')')) or
                '_' in placeholder  # Underscores
            )
            is_blank_field = placeholder.endswith(': ') or placeholder.endswith(':')

            # Build patterns
            patterns_to_try = [placeholder]
            if is_blank_field:
                base = placeholder.rstrip(': \t')
                patterns_to_try.extend([
                    base + ':\t',
                    base + ':  ',
                    base + ': ',
                    base + ': _____',
                ])
            
            # Collect all occurrences
            occurrences = []
            
            for para in self.doc.paragraphs:
                full_text = ''.join([run.text for run in para.runs])
                for pattern in patterns_to_try:
                    if pattern in full_text:
                        occurrences.append((para, pattern))
                        break
            
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_text = ''.join([run.text for run in para.runs])
                            for pattern in patterns_to_try:
                                if pattern in full_text:
                                    occurrences.append((para, pattern))
                                    break
            
            # Get target occurrence
            if position_index >= len(occurrences):
                return False
            
            target_para, matching_pattern = occurrences[position_index]
            full_para_text = ''.join([run.text for run in target_para.runs])
            
            # SIMPLE REPLACEMENT
            if is_explicit_placeholder:
                new_text = full_para_text.replace(matching_pattern, value, 1)
            else:
                # Blank field: keep label
                label_pos = full_para_text.find(matching_pattern)
                if label_pos != -1:
                    new_text = full_para_text[:label_pos + len(matching_pattern)] + value
                else:
                    return False
            
            # Write back
            for run in target_para.runs:
                r = run._element
                r.getparent().remove(r)
            
            target_para.add_run(new_text)
            return True
        except Exception as e:
            print(f"Error replacing placeholder at position: {e}")
            return False
    
    def replace_multiple_placeholders(self, replacements: Dict[str, str]) -> Dict[str, bool]:
        """
        Replace multiple placeholders at once
        
        Args:
            replacements: Dictionary mapping placeholder -> value
        
        Returns:
            Dictionary mapping placeholder -> success status
        """
        results = {}
        for placeholder, value in replacements.items():
            results[placeholder] = self.replace_placeholder(placeholder, value)
        return results
    
    def save_document(self, output_path: str) -> bool:
        """Save the modified document to a new file"""
        try:
            self.doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False

    def replace_placeholder_in_section(self, placeholder: str, value: str, section_keyword: Optional[str] = None) -> bool:
        """
        Replace a placeholder only in a specific section of the document.
        Counts occurrences within the section to find the right one.
        
        Args:
            placeholder: The placeholder text to find (e.g., "Address: ")
            value: The replacement value
            section_keyword: Section identifier (e.g., "investor", "company")
        
        Returns:
            True if replacement was successful
        """
        try:
            # Detect if this is a blank field placeholder
            is_blank_field = (
                (': ' in placeholder and placeholder.endswith(': ')) or
                placeholder.endswith(':')
            )
            
            if not section_keyword:
                # No section keyword, just use regular replacement
                return self.replace_placeholder(placeholder, value)
            
            # Find the LAST section header with this keyword
            # (section headers are at the END of the document, not scattered throughout)
            section_start_para = None
            for i in range(len(self.doc.paragraphs) - 1, -1, -1):  # Search backwards!
                para_text = self.doc.paragraphs[i].text.upper()
                # Look for isolated keyword or with colons (like "INVESTOR:")
                if (section_keyword.upper() == para_text.strip() or 
                    section_keyword.upper() + ':' in para_text or
                    para_text.startswith(section_keyword.upper() + ':')):
                    section_start_para = i
                    break
            
            if section_start_para is None:
                # Section not found, use regular replacement
                return self.replace_placeholder(placeholder, value)
            
            # Find first occurrence of placeholder AFTER section start
            for i in range(section_start_para, len(self.doc.paragraphs)):
                para = self.doc.paragraphs[i]
                full_para_text = ''.join([run.text for run in para.runs])
                
                if placeholder in full_para_text:
                    # Found it! Now replace
                    if is_blank_field:
                        # For blank fields: keep label, replace blank part
                        label_pos = full_para_text.find(placeholder)
                        if label_pos != -1:
                            new_text = full_para_text[:label_pos + len(placeholder)] + value
                        else:
                            continue
                    else:
                        # For explicit placeholders: replace entire placeholder
                        new_text = full_para_text.replace(placeholder, value, 1)
                    
                    # Write back
                    for run in para.runs:
                        r = run._element
                        r.getparent().remove(r)
                    
                    para.add_run(new_text)
                    return True
            
            # Also check in tables
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_para_text = ''.join([run.text for run in para.runs])
                            
                            if placeholder in full_para_text:
                                if is_blank_field:
                                    label_pos = full_para_text.find(placeholder)
                                    if label_pos != -1:
                                        new_text = full_para_text[:label_pos + len(placeholder)] + value
                                    else:
                                        continue
                                else:
                                    new_text = full_para_text.replace(placeholder, value, 1)
                                
                                for run in para.runs:
                                    r = run._element
                                    r.getparent().remove(r)
                                
                                para.add_run(new_text)
                                return True
            
            return False
        except Exception as e:
            print(f"Error in section-aware replacement: {e}")
            return False
