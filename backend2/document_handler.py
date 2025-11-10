"""
Document Handler for backend2
Handles loading, modifying, and saving Word documents using python-docx
"""

import os
import sys
import re
from docx import Document
from typing import Dict, List, Tuple, Optional


class DocumentHandler:
    def __init__(self, doc_path: str):
        """Initialize document handler with path to .docx file"""
        self.doc_path = doc_path
        self.doc = None
        self.full_text = ""
        
    def load_document(self) -> bool:
        """Load the .docx document"""
        try:
            self.doc = Document(self.doc_path)
            self._extract_text_structure()
            return True
        except Exception as e:
            print(f"Error loading document: {e}", file=sys.stderr)
            return False
    
    def _extract_text_structure(self):
        """Extract text while preserving structure"""
        self.full_text = ""
        
        # Extract from regular paragraphs
        for para in self.doc.paragraphs:
            para_text = para.text
            self.full_text += para_text + "\n"
        
        # Extract from table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text.strip():
                        self.full_text += cell_text + "\n"
    
    def get_full_text(self) -> str:
        """Return extracted full text"""
        return self.full_text
    
    def _replace_text_preserving_format(self, para, new_text: str, label_start_pos: Optional[int] = None):
        """
        Replace text in paragraph while preserving formatting character-by-character.
        Maps each character in new_text to corresponding character in old text to preserve formatting.
        
        Args:
            para: The paragraph to modify
            new_text: The new text to replace with
            label_start_pos: Optional position where label starts (for label fields, preserve label formatting)
        """
        # Build character-to-run mapping to preserve formatting BEFORE clearing runs
        char_to_run = []
        char_pos = 0
        runs_list = list(para.runs)  # Copy list before clearing
        old_text = ''.join([run.text for run in runs_list])
        
        for run in runs_list:
            run_text = run.text
            for i in range(len(run_text)):
                char_to_run.append((char_pos + i, run))
            char_pos += len(run_text)
        
        # Clear all runs
        for run in para.runs:
            r = run._element
            r.getparent().remove(r)
        
        # Map new text characters to old text positions to preserve formatting
        # We need to find where the replacement happened and map accordingly
        
        if not char_to_run:
            # No old text, just add plain text
            para.add_run(new_text)
            return
        
        # Find the difference between old and new text to locate replacement
        # Simple approach: find longest common prefix and suffix
        prefix_len = 0
        suffix_len = 0
        
        # Find common prefix
        for i in range(min(len(old_text), len(new_text))):
            if old_text[i] == new_text[i]:
                prefix_len += 1
            else:
                break
        
        # Find common suffix
        old_rev = old_text[::-1]
        new_rev = new_text[::-1]
        for i in range(min(len(old_text) - prefix_len, len(new_text) - prefix_len)):
            if old_rev[i] == new_rev[i]:
                suffix_len += 1
            else:
                break
        
        # Build runs for new text, preserving formatting from corresponding positions in old text
        current_run = None
        current_run_format = None
        
        for i, char in enumerate(new_text):
            # Determine which run format to use for this character
            target_run = None
            old_pos = None
            
            if i < prefix_len:
                # Before replacement - use original position
                old_pos = i
                target_run = char_to_run[i][1]
            elif i >= len(new_text) - suffix_len:
                # After replacement - map to corresponding position in old text
                # Calculate offset: new_suffix_start - old_suffix_start
                new_suffix_start = len(new_text) - suffix_len
                old_suffix_start = len(old_text) - suffix_len
                old_pos = old_suffix_start + (i - new_suffix_start)
                if old_pos < len(char_to_run):
                    target_run = char_to_run[old_pos][1]
            else:
                # In replacement region - use formatting from before replacement (or after if at start)
                if prefix_len > 0 and prefix_len < len(char_to_run):
                    # Use formatting from character just before replacement
                    target_run = char_to_run[prefix_len - 1][1]
                elif suffix_len > 0 and (len(old_text) - suffix_len) < len(char_to_run):
                    # Use formatting from character just after replacement
                    target_run = char_to_run[len(old_text) - suffix_len][1]
                elif label_start_pos is not None and label_start_pos < len(char_to_run):
                    # For label fields, use label's formatting
                    target_run = char_to_run[label_start_pos][1]
                elif len(char_to_run) > 0:
                    # Fallback: use formatting from last character
                    target_run = char_to_run[-1][1]
            
            # Check if we need to start a new run (formatting changed)
            if target_run:
                # Get formatting properties for this character
                char_format = {
                    'bold': target_run.bold,
                    'italic': target_run.italic,
                    'underline': target_run.underline,
                    'font_name': target_run.font.name if target_run.font.name else None,
                    'font_size': target_run.font.size if target_run.font.size else None,
                    'font_color': target_run.font.color.rgb if target_run.font.color.rgb else None,
                }
                
                # Compare formatting (handle None values properly)
                formatting_changed = True
                if current_run_format is not None:
                    formatting_changed = (
                        char_format['bold'] != current_run_format['bold'] or
                        char_format['italic'] != current_run_format['italic'] or
                        char_format['underline'] != current_run_format['underline'] or
                        char_format['font_name'] != current_run_format['font_name'] or
                        char_format['font_size'] != current_run_format['font_size'] or
                        char_format['font_color'] != current_run_format['font_color']
                    )
                
                # Start new run if formatting changed
                if formatting_changed:
                    current_run = para.add_run(char)
                    current_run_format = char_format
                    
                    # Apply formatting
                    if char_format['bold'] is not None:
                        current_run.bold = char_format['bold']
                    if char_format['italic'] is not None:
                        current_run.italic = char_format['italic']
                    if char_format['underline'] is not None:
                        current_run.underline = char_format['underline']
                    if char_format['font_name']:
                        current_run.font.name = char_format['font_name']
                    if char_format['font_size']:
                        current_run.font.size = char_format['font_size']
                    if char_format['font_color']:
                        current_run.font.color.rgb = char_format['font_color']
                else:
                    # Continue current run
                    current_run.text += char
            else:
                # No formatting info, add to current run or create plain run
                if current_run is None:
                    current_run = para.add_run(char)
                    current_run_format = None
                else:
                    current_run.text += char
    
    def replace_placeholder(self, placeholder: str, value: str) -> bool:
        """
        Replace placeholder with value.
        
        Logic:
        - Explicit placeholder ([text], {text}, (text)): Replace ENTIRE placeholder with value
        - Label field (text like "Date", "No.", "The Sum of"): Keep label, add space, then insert value
        
        Args:
            placeholder: The placeholder text to find
            value: The replacement value
        
        Returns:
            True if replacement was successful
        """
        try:
            replaced_count = 0
            
            # Determine type: explicit placeholder or label field
            is_explicit_placeholder = (
                (placeholder.startswith('[') and placeholder.endswith(']')) or
                (placeholder.startswith('{') and placeholder.endswith('}')) or
                (placeholder.startswith('(') and placeholder.endswith(')')) or
                (placeholder.startswith('<') and placeholder.endswith('>')) or
                '_' in placeholder  # Underscores are explicit
            )
            
            is_label_field = not is_explicit_placeholder  # Any non-bracketed placeholder is a label field
            
            # Build list of patterns to try (handle whitespace variations)
            patterns_to_try = [placeholder]
            if is_label_field:
                # For label fields, try variations with/without colon, spaces, etc.
                base = placeholder.rstrip(': \t')  # Get the label without trailing space/colon
                patterns_to_try.extend([
                    base + ':\t',    # Tab variant
                    base + ':  ',    # Double space variant
                    base + ': ',     # Space variant
                    base + ':',     # Just colon
                    base,           # Just the label name (for cases like "Date2023-10-01" or "The Sum of1200.00")
                ])
            
            # Replace in paragraphs
            for para in self.doc.paragraphs:
                full_para_text = ''.join([run.text for run in para.runs])
                
                for pattern in patterns_to_try:
                    if pattern in full_para_text:
                        if is_explicit_placeholder:
                            # Replace entire placeholder
                            new_text = full_para_text.replace(pattern, value, 1)
                        else:
                            # Label field: keep label, add space, then insert value
                            label_pos = full_para_text.find(pattern)
                            if label_pos != -1:
                                label_end = label_pos + len(pattern)
                                remaining_text = full_para_text[label_end:]
                                
                                # Strip trailing spaces from pattern to get actual label end
                                label_without_trailing_space = pattern.rstrip(' \t')
                                actual_label_end = label_pos + len(label_without_trailing_space)
                                remaining_after_label = full_para_text[actual_label_end:]
                                
                                # Check what comes after the label
                                if remaining_text and not remaining_text[0].isspace():
                                    # There's text immediately after label (no space), replace it
                                    # Find where the existing value ends (look for space, newline, or end)
                                    match = re.match(r'^([^\s\n]+)', remaining_text)
                                    if match:
                                        # Replace the existing value
                                        existing_value_end = match.end()
                                        new_text = full_para_text[:label_end] + ' ' + value + remaining_text[existing_value_end:]
                                    else:
                                        # No clear existing value, just append
                                        new_text = full_para_text[:label_end] + ' ' + value
                                else:
                                    # There's whitespace/blank lines after label - REPLACE them with value
                                    # For label fields, we want: label + ' ' + value (all on same line)
                                    # Replace ALL whitespace/newlines after label with just space + value
                                    match = re.match(r'^[\s\n\t]+', remaining_after_label)
                                    if match:
                                        # Replace the blank content
                                        after_whitespace = remaining_after_label[match.end():]
                                        if after_whitespace.strip():
                                            # There's content after whitespace, keep it
                                            new_text = full_para_text[:actual_label_end] + ' ' + value + ' ' + after_whitespace
                                        else:
                                            # No content after whitespace, just replace with label + space + value
                                            new_text = full_para_text[:actual_label_end] + ' ' + value
                                    else:
                                        # No blank content, just append value with space
                                        new_text = full_para_text[:actual_label_end] + ' ' + value
                            else:
                                continue
                        
                        if new_text != full_para_text:
                            # Preserve formatting by modifying runs in place
                            self._replace_text_preserving_format(para, new_text, label_pos if not is_explicit_placeholder else None)
                            replaced_count += 1
                            break  # Move to next paragraph
            
            # Replace in table cells
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_para_text = ''.join([run.text for run in para.runs])
                            
                            for pattern in patterns_to_try:
                                if pattern in full_para_text:
                                    if is_explicit_placeholder:
                                        # Replace only the FIRST occurrence (same placeholder might mean different things)
                                        new_text = full_para_text.replace(pattern, value, 1)
                                    else:
                                        # Label field: keep label, add space, then insert value
                                        label_pos = full_para_text.find(pattern)
                                        if label_pos != -1:
                                            label_end = label_pos + len(pattern)
                                            remaining_text = full_para_text[label_end:]
                                            
                                            # Strip trailing spaces from pattern to get actual label end
                                            label_without_trailing_space = pattern.rstrip(' \t')
                                            actual_label_end = label_pos + len(label_without_trailing_space)
                                            remaining_after_label = full_para_text[actual_label_end:]
                                            
                                            # Check what comes after the label
                                            if remaining_text and not remaining_text[0].isspace():
                                                # There's text immediately after label, replace it
                                                match = re.match(r'^([^\s\n]+)', remaining_text)
                                                if match:
                                                    existing_value_end = match.end()
                                                    new_text = full_para_text[:label_end] + ' ' + value + remaining_text[existing_value_end:]
                                                else:
                                                    new_text = full_para_text[:label_end] + ' ' + value
                                            else:
                                                # There's whitespace/blank lines after label - REPLACE them with value
                                                match = re.match(r'^[\s\n\t]+', remaining_after_label)
                                                if match:
                                                    after_whitespace = remaining_after_label[match.end():]
                                                    if after_whitespace.strip():
                                                        new_text = full_para_text[:actual_label_end] + ' ' + value + ' ' + after_whitespace
                                                    else:
                                                        new_text = full_para_text[:actual_label_end] + ' ' + value
                                                else:
                                                    new_text = full_para_text[:actual_label_end] + ' ' + value
                                        else:
                                            continue
                                    
                                    if new_text != full_para_text:
                                        # Preserve formatting by modifying runs in place
                                        self._replace_text_preserving_format(para, new_text, label_pos if not is_explicit_placeholder else None)
                                        replaced_count += 1
                                        break  # Move to next paragraph
            
            return replaced_count > 0
        except Exception as e:
            print(f"Error replacing placeholder: {e}", file=sys.stderr)
            return False
    
    def replace_placeholder_at_position(self, placeholder: str, value: str, position_index: int = 0) -> bool:
        """
        Replace a specific occurrence (by position) of a placeholder.
        
        Args:
            placeholder: The placeholder text to find
            value: The replacement value  
            position_index: Which occurrence (0=first, 1=second, etc.)
        
        Returns:
            True if replacement was successful
        """
        try:
            # Determine type
            is_explicit_placeholder = (
                (placeholder.startswith('[') and placeholder.endswith(']')) or
                (placeholder.startswith('{') and placeholder.endswith('}')) or
                (placeholder.startswith('(') and placeholder.endswith(')')) or
                (placeholder.startswith('<') and placeholder.endswith('>')) or
                '_' in placeholder
            )
            is_label_field = not is_explicit_placeholder
            
            # Build patterns
            patterns_to_try = [placeholder]
            if is_label_field:
                base = placeholder.rstrip(': \t\n').strip()  # Strip both leading and trailing whitespace
                # Try common variations - use a reasonable set of leading space patterns
                # Most documents use 0-20 spaces for indentation
                for i in range(0, 21):  # 0 to 20 leading spaces
                    leading_spaces = ' ' * i
                    patterns_to_try.extend([
                        leading_spaces + base + ':\t',
                        leading_spaces + base + ':  ',
                        leading_spaces + base + ': ',
                        leading_spaces + base + ':',
                    ])
                # Also try without leading spaces (already covered but explicit)
                patterns_to_try.extend([
                    base + ':\t',
                    base + ':  ',
                    base + ': ',
                    base + ':',
                    base,           # Just the label name
                ])
            
            # Collect all occurrences
            occurrences = []
            seen_paragraphs = set()  # Track which paragraphs we've already added
            
            # For label fields, use normalized matching to handle whitespace variations
            if is_label_field:
                base_label = placeholder.rstrip(': \t\n').strip().lower()
                
                def matches_label(text, label_base):
                    """Check if text contains the label (ignoring leading/trailing whitespace)"""
                    # Normalize: remove leading/trailing whitespace and check if label: appears
                    normalized = text.strip().lower()
                    # Check if label: appears anywhere (with flexible whitespace)
                    import re
                    pattern = r'\s*' + re.escape(label_base) + r'\s*:'
                    return bool(re.search(pattern, normalized))
                
                def extract_label_pattern(text, label_base):
                    """Extract the actual label pattern from text (e.g., 'Address:', ' Address: ', etc.)"""
                    import re
                    # Find the label with flexible whitespace - preserve ALL whitespace (spaces, tabs, newlines)
                    # Match: any whitespace + label + any whitespace + colon + any whitespace
                    pattern = r'(\s*' + re.escape(label_base) + r'\s*:\s*)'
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        return match.group(1)  # Return the actual matched text WITH all whitespace preserved
                    return None
                
                for para in self.doc.paragraphs:
                    full_text = ''.join([run.text for run in para.runs])
                    para_id = id(para)
                    if para_id not in seen_paragraphs and matches_label(full_text, base_label):
                        # Extract the actual pattern from the text (handles any whitespace variation)
                        matching_pattern = extract_label_pattern(full_text, base_label)
                        if not matching_pattern:
                            # Fallback: try patterns_to_try
                            for pattern in patterns_to_try:
                                if pattern in full_text:
                                    matching_pattern = pattern
                                    break
                        
                        if matching_pattern:
                            occurrences.append((para, matching_pattern, full_text))
                            seen_paragraphs.add(para_id)
                
                for table in self.doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full_text = ''.join([run.text for run in para.runs])
                                para_id = id(para)
                                if para_id not in seen_paragraphs and matches_label(full_text, base_label):
                                    # Extract the actual pattern from the text
                                    matching_pattern = extract_label_pattern(full_text, base_label)
                                    if not matching_pattern:
                                        # Fallback: try patterns_to_try
                                        for pattern in patterns_to_try:
                                            if pattern in full_text:
                                                matching_pattern = pattern
                                                break
                                    
                                    if matching_pattern:
                                        occurrences.append((para, matching_pattern, full_text))
                                        seen_paragraphs.add(para_id)
            else:
                # For explicit placeholders, use exact matching
                for para in self.doc.paragraphs:
                    full_text = ''.join([run.text for run in para.runs])
                    for pattern in patterns_to_try:
                        if pattern in full_text:
                            para_id = id(para)
                            if para_id not in seen_paragraphs:
                                occurrences.append((para, pattern, full_text))
                                seen_paragraphs.add(para_id)
                            break
                
                for table in self.doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full_text = ''.join([run.text for run in para.runs])
                                para_id = id(para)
                                if para_id in seen_paragraphs:
                                    continue
                                for pattern in patterns_to_try:
                                    if pattern in full_text:
                                        occurrences.append((para, pattern, full_text))
                                        seen_paragraphs.add(para_id)
                                        break
            
            # Get target occurrence
            if position_index >= len(occurrences):
                return False
            
            target_para, matching_pattern, full_para_text = occurrences[position_index]
            
            # Replace
            if is_explicit_placeholder:
                new_text = full_para_text.replace(matching_pattern, value, 1)
            else:
                # Label field: keep label, add space, then insert value
                label_pos = full_para_text.find(matching_pattern)
                if label_pos != -1:
                    label_end = label_pos + len(matching_pattern)
                    remaining_text = full_para_text[label_end:]
                    
                    # Strip trailing spaces from pattern to get actual label end
                    label_without_trailing_space = matching_pattern.rstrip(' \t')
                    actual_label_end = label_pos + len(label_without_trailing_space)
                    remaining_after_label = full_para_text[actual_label_end:]
                    
                    # Check what comes after the label
                    if remaining_text and not remaining_text[0].isspace():
                        # There's text immediately after label, replace it
                        match = re.match(r'^([^\s\n]+)', remaining_text)
                        if match:
                            existing_value_end = match.end()
                            new_text = full_para_text[:label_end] + ' ' + value + remaining_text[existing_value_end:]
                        else:
                            new_text = full_para_text[:label_end] + ' ' + value
                    else:
                        # There's whitespace/blank lines after label - REPLACE them with value
                        match = re.match(r'^[\s\n\t]+', remaining_after_label)
                        if match:
                            after_whitespace = remaining_after_label[match.end():]
                            if after_whitespace.strip():
                                new_text = full_para_text[:actual_label_end] + ' ' + value + ' ' + after_whitespace
                            else:
                                new_text = full_para_text[:actual_label_end] + ' ' + value
                        else:
                            new_text = full_para_text[:actual_label_end] + ' ' + value
                else:
                    return False
            
            # Write back with formatting preservation
            self._replace_text_preserving_format(target_para, new_text, label_pos if not is_explicit_placeholder else None)
            return True
        except Exception as e:
            print(f"Error replacing placeholder at position: {e}", file=sys.stderr)
            return False
    
    def save_document(self, output_path: str) -> bool:
        """Save the modified document to a new file"""
        try:
            self.doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}", file=sys.stderr)
            return False

